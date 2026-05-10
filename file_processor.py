"""
معالجة الملفات وبناء قاعدة المعرفة باستخدام FAISS
"""

import os
import json
import hashlib
import pickle
from typing import List, Optional, Tuple
from pathlib import Path

import PyPDF2
import docx
import pandas as pd
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain_core.documents import Document
from openai import OpenAI

from config import (
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    VECTOR_DB_PATH,
    UPLOADS_DIR,
    SUPPORTED_FILE_TYPES,
)


def extract_text_from_pdf(file_path: str) -> str:
    """استخراج النص من ملف PDF"""
    text = ""
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    except Exception as e:
        text = f"خطأ في قراءة ملف PDF: {str(e)}"
    return text


def extract_text_from_docx(file_path: str) -> str:
    """استخراج النص من ملف Word"""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
    except Exception as e:
        text = f"خطأ في قراءة ملف Word: {str(e)}"
    return text


def extract_text_from_csv(file_path: str) -> str:
    """استخراج النص من ملف CSV"""
    try:
        df = pd.read_csv(file_path, encoding="utf-8")
        return df.to_string(index=False)
    except Exception:
        try:
            df = pd.read_csv(file_path, encoding="latin-1")
            return df.to_string(index=False)
        except Exception as e:
            return f"خطأ في قراءة ملف CSV: {str(e)}"


def extract_text_from_txt(file_path: str) -> str:
    """استخراج النص من ملف نصي"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        try:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()
        except Exception as e:
            return f"خطأ في قراءة الملف النصي: {str(e)}"


def extract_text_from_file(file_path: str) -> str:
    """استخراج النص من أي ملف مدعوم"""
    ext = Path(file_path).suffix.lower().lstrip(".")
    extractors = {
        "pdf": extract_text_from_pdf,
        "docx": extract_text_from_docx,
        "doc": extract_text_from_docx,
        "csv": extract_text_from_csv,
        "txt": extract_text_from_txt,
        "md": extract_text_from_txt,
    }
    extractor = extractors.get(ext)
    if extractor:
        return extractor(file_path)
    return f"نوع الملف '{ext}' غير مدعوم."


def get_file_hash(file_path: str) -> str:
    """حساب hash للملف للتحقق من التغييرات"""
    hasher = hashlib.md5()
    with open(file_path, "rb") as f:
        buf = f.read(65536)
        while len(buf) > 0:
            hasher.update(buf)
            buf = f.read(65536)
    return hasher.hexdigest()


class KnowledgeBase:
    """قاعدة المعرفة المبنية على FAISS"""

    def __init__(self, api_key: str, session_id: str = "default"):
        self.api_key = api_key
        self.session_id = session_id
        self.db_path = os.path.join(VECTOR_DB_PATH, session_id)
        self.vectorstore: Optional[FAISS] = None
        self.documents_info: List[dict] = []
        self.embeddings = None
        self._init_embeddings()
        self._load_existing()

    def _init_embeddings(self):
        """تهيئة نموذج التضمين"""
        try:
            self.embeddings = OpenAIEmbeddings(
                api_key=self.api_key,
                model="text-embedding-3-small",
            )
        except Exception:
            self.embeddings = None

    def _load_existing(self):
        """تحميل قاعدة المعرفة الموجودة"""
        info_path = os.path.join(self.db_path, "docs_info.json")
        if os.path.exists(info_path):
            try:
                with open(info_path, "r", encoding="utf-8") as f:
                    self.documents_info = json.load(f)
            except Exception:
                self.documents_info = []

        if self.embeddings and os.path.exists(self.db_path):
            try:
                self.vectorstore = FAISS.load_local(
                    self.db_path,
                    self.embeddings,
                    allow_dangerous_deserialization=True,
                )
            except Exception:
                self.vectorstore = None

    def add_file(self, file_path: str, file_name: str) -> Tuple[bool, str]:
        """إضافة ملف إلى قاعدة المعرفة"""
        if not self.embeddings:
            return False, "لم يتم تهيئة نموذج التضمين. تأكد من صحة مفتاح API."

        try:
            text = extract_text_from_file(file_path)
            if not text or len(text.strip()) < 10:
                return False, "لم يتم استخراج نص كافٍ من الملف."

            splitter = RecursiveCharacterTextSplitter(
                chunk_size=CHUNK_SIZE,
                chunk_overlap=CHUNK_OVERLAP,
                separators=["\n\n", "\n", ".", "،", " ", ""],
            )
            chunks = splitter.split_text(text)
            docs = [
                Document(
                    page_content=chunk,
                    metadata={"source": file_name, "chunk_id": i},
                )
                for i, chunk in enumerate(chunks)
            ]

            if self.vectorstore is None:
                self.vectorstore = FAISS.from_documents(docs, self.embeddings)
            else:
                self.vectorstore.add_documents(docs)

            os.makedirs(self.db_path, exist_ok=True)
            self.vectorstore.save_local(self.db_path)

            file_hash = get_file_hash(file_path)
            self.documents_info.append({
                "name": file_name,
                "path": file_path,
                "hash": file_hash,
                "chunks": len(chunks),
                "chars": len(text),
            })
            info_path = os.path.join(self.db_path, "docs_info.json")
            with open(info_path, "w", encoding="utf-8") as f:
                json.dump(self.documents_info, f, ensure_ascii=False, indent=2)

            return True, f"تم إضافة الملف بنجاح ({len(chunks)} قطعة نصية)"

        except Exception as e:
            return False, f"خطأ في معالجة الملف: {str(e)}"

    def search(self, query: str, k: int = 4) -> List[Document]:
        """البحث في قاعدة المعرفة"""
        if self.vectorstore is None:
            return []
        try:
            return self.vectorstore.similarity_search(query, k=k)
        except Exception:
            return []

    def clear(self):
        """مسح قاعدة المعرفة"""
        self.vectorstore = None
        self.documents_info = []
        import shutil
        if os.path.exists(self.db_path):
            shutil.rmtree(self.db_path)

    def get_documents_list(self) -> List[str]:
        """الحصول على قائمة الملفات المضافة"""
        return [doc["name"] for doc in self.documents_info]

    def has_documents(self) -> bool:
        """التحقق من وجود وثائق"""
        return self.vectorstore is not None and len(self.documents_info) > 0
